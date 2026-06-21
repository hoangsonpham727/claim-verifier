from pathlib import Path

from docx import Document
from docx.shared import Pt
import pymupdf

# Sample documents are a shared, language-neutral asset (../shared/samples),
# resolved relative to this file so the script works from any working directory.
SAMPLES = Path(__file__).resolve().parent.parent / "shared" / "samples"
SAMPLES.mkdir(parents=True, exist_ok=True)

# ── 1) Main Word document the user opens in Word ──────────────────────────────
doc = Document()
doc.add_heading("MUTUAL NON-DISCLOSURE AGREEMENT", level=0)

sections = [
    ("Section 1. Parties",
     "This Agreement is entered into between Acme Corporation and Beta Industries LLC as of 1 March 2024."),
    ("Section 2. Definitions",
     "Confidential Information means any non-public business, technical, or financial information disclosed by one party to the other party."),
    ("Section 3. Confidentiality Obligation",
     "The Receiving Party shall keep all Confidential Information strictly secret and shall not disclose it to any third party without prior written consent. "
     # demo claim → SUPPORTED by the Data Processing Addendum (Art. 2)
     "Pursuant to the Data Processing Addendum, the Receiving Party shall not "
     "disclose Confidential Information to any third party without the prior "
     "written consent of the Disclosing Party."),
    ("Section 4. Term",
     "The confidentiality obligations shall remain in effect for a period of five (5) years from the date of disclosure."),
    ("Section 5. Survival",
     "As provided in Section 4, the duty to protect Confidential Information continues for five years following the date of disclosure."),
    ("Section 6. Incorporated Terms",
     "The definition of Confidential Information also includes the categories set forth in the Master Services Agreement dated 1 January 2023. "
     # demo claim → CONTRADICTED by the Master Services Agreement (Art. 3 says
     # IP remains with the Disclosing Party, not the Receiving Party)
     "As set forth in the Master Services Agreement, all intellectual property "
     "created under this Agreement becomes the property of the Receiving Party."),
    ("Section 7. Return of Materials",
     "Upon termination of this Agreement, the Receiving Party shall return or destroy all Confidential Information within thirty (30) days."),
    ("Section 8. Indemnification",
     "Each party shall indemnify the other party for any losses arising from a material breach of this Agreement. "
     # demo claim → UNADDRESSED (no source mentions insurance at all)
     "As provided in the Master Services Agreement, each party shall maintain "
     "cyber-liability insurance of at least five million dollars ($5,000,000)."),
    ("Section 9. Governing Law",
     "This Agreement shall be governed by and construed in accordance with the laws of the State of Delaware."),
    # ── Demo cite-check representations ───────────────────────────────────────
    # Three sentences, each carrying a citation marker so segment.py flags them as
    # claims, engineered to land on a distinct verdict against the sources:
    #   • supported    — matches the Data Processing Addendum verbatim-ish
    #   • contradicted — the DPA says the opposite (EU-only vs United States)
    #   • unaddressed  — no source mentions uptime / SLAs at all
    ("Section 10. Data Protection Representations.",
     "Pursuant to the Data Processing Addendum, the Receiving Party shall notify "
     "the Disclosing Party of any personal data breach within seventy-two (72) "
     "hours of discovery. "
     "As set forth in the Data Processing Addendum, the Receiving Party may "
     "disclose Confidential Information to any third party without the prior "
     "written consent of the Disclosing Party. "
     "As provided in the Data Processing Addendum, the Provider guarantees 99.9% "
     "monthly service uptime availability."),
]
for heading, body in sections:
    doc.add_heading(heading, level=2)
    doc.add_paragraph(body)

doc.save(str(SAMPLES / "NDA-2024.docx"))
print(f"wrote {SAMPLES / 'NDA-2024.docx'}")

# ── 2) External source file (PDF) the user uploads ────────────────────────────
pdf = pymupdf.open()
page = pdf.new_page()
lines = [
    "MASTER SERVICES AGREEMENT",
    "Dated 1 January 2023.",
    "",
    "Article 1. Definitions.",
    "For the purposes of this Agreement, Confidential Information includes all source",
    "code, product designs, customer lists, and pricing information disclosed by either party.",
    "",
    "Article 2. Protection Period.",
    "Each party shall protect Confidential Information for a period of ten (10) years from",
    "the date of disclosure.",
    "",
    "Article 3. Ownership.",
    "All intellectual property created under this Agreement remains the property of the",
    "Disclosing Party.",
    "",
    "Article 4. Governing Law.",
    "This Agreement shall be governed by the laws of the State of New York.",
]
y = 72
for ln in lines:
    page.insert_text((72, y), ln, fontsize=11)
    y += 20
pdf.save(str(SAMPLES / "Master-Services-Agreement.pdf"))
print(f"wrote {SAMPLES / 'Master-Services-Agreement.pdf'}")

# ── 3) Second external source (PDF) — Data Processing Addendum ─────────────────
# Covers topics NOT in the NDA or MSA, so the demo claims resolve unambiguously:
#   breach notification (72h)  → supports the "72 hours" claim
#   data residency (EU only)   → contradicts the "United States" claim
#   (nothing about uptime/SLA) → leaves the "99.9% uptime" claim unaddressed
dpa = pymupdf.open()
dpa_page = dpa.new_page()
dpa_lines = [
    "DATA PROCESSING ADDENDUM",
    "Dated 1 February 2024.",
    "",
    "Article 1. Breach Notification.",
    "The Receiving Party shall notify the Disclosing Party of any personal data",
    "breach within seventy-two (72) hours of discovery.",
    "",
    "Article 2. Non-Disclosure.",
    "The Receiving Party shall keep all Confidential Information strictly secret and",
    "shall not disclose it to any third party without the prior written consent of",
    "the Disclosing Party.",
    "",
    "Article 3. Audit Rights.",
    "The Disclosing Party may audit the Receiving Party's data-protection",
    "practices once per calendar year.",
]
y = 72
for ln in dpa_lines:
    dpa_page.insert_text((72, y), ln, fontsize=11)
    y += 20
dpa.save(str(SAMPLES / "Data-Processing-Addendum.pdf"))
print(f"wrote {SAMPLES / 'Data-Processing-Addendum.pdf'}")
