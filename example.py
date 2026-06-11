from docx import Document
from docx.shared import Pt
import pymupdf

# ── 1) Main Word document the user opens in Word ──────────────────────────────
doc = Document()
doc.add_heading("MUTUAL NON-DISCLOSURE AGREEMENT", level=0)

sections = [
    ("Section 1. Parties",
     "This Agreement is entered into between Acme Corporation and Beta Industries LLC as of 1 March 2024."),
    ("Section 2. Definitions",
     "Confidential Information means any non-public business, technical, or financial information disclosed by one party to the other party."),
    ("Section 3. Confidentiality Obligation",
     "The Receiving Party shall keep all Confidential Information strictly secret and shall not disclose it to any third party without prior written consent."),
    ("Section 4. Term",
     "The confidentiality obligations shall remain in effect for a period of five (5) years from the date of disclosure."),
    ("Section 5. Survival",
     "As provided in Section 4, the duty to protect Confidential Information continues for five years following the date of disclosure."),
    ("Section 6. Incorporated Terms",
     "The definition of Confidential Information also includes the categories set forth in the Master Services Agreement dated 1 January 2023."),
    ("Section 7. Return of Materials",
     "Upon termination of this Agreement, the Receiving Party shall return or destroy all Confidential Information within thirty (30) days."),
    ("Section 8. Indemnification",
     "Each party shall indemnify the other party for any losses arising from a material breach of this Agreement."),
    ("Section 9. Governing Law",
     "This Agreement shall be governed by and construed in accordance with the laws of the State of Delaware."),
]
for heading, body in sections:
    doc.add_heading(heading, level=2)
    doc.add_paragraph(body)

doc.save("samples/NDA-2024.docx")
print("wrote samples/NDA-2024.docx")

# ── 2) External source file (PDF) the user uploads ────────────────────────────
pdf = pymupdf.open()
page = pdf.new_page()
lines = [
    "MASTER SERVICES AGREEMENT",
    "Dated 1 January 2023.",
    "",
    "Article 1. Definitions.",
    "For the purposes of this Agreement, Confidential Information includes all source",
    "code, product designs, customer lists, and pricing information disclosed by either party.",
    "",
    "Article 2. Protection Period.",
    "Each party shall protect Confidential Information for a period of ten (10) years from",
    "the date of disclosure.",
    "",
    "Article 3. Ownership.",
    "All intellectual property created under this Agreement remains the property of the",
    "Disclosing Party.",
    "",
    "Article 4. Governing Law.",
    "This Agreement shall be governed by the laws of the State of New York.",
]
y = 72
for ln in lines:
    page.insert_text((72, y), ln, fontsize=11)
    y += 20
pdf.save("samples/Master-Services-Agreement.pdf")
print("wrote samples/Master-Services-Agreement.pdf")
